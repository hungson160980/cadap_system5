import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
import io
import re
from datetime import datetime
import json
import time

# Import có điều kiện
try:
    import plotly.express as px
    import plotly.graph_objects as go
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

try:
    import google.generativeai as genai
    GENAI_AVAILABLE = True
except ImportError:
    GENAI_AVAILABLE = False

# Cấu hình trang
st.set_page_config(
    page_title="Hệ Thống Thẩm Định Phương Án Kinh Doanh",
    page_icon="🏦",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS tùy chỉnh - Cải thiện giao diện tabs
st.markdown("""
<style>
    /* Header chính */
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        padding: 1rem 0;
        background: linear-gradient(90deg, #e3f2fd 0%, #bbdefb 100%);
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    
    /* Metric cards */
    .metric-card {
        background: white;
        padding: 1rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    /* Tối ưu hóa Tabs */
    .stTabs {
        background-color: #f8f9fa;
        padding: 10px;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        overflow-x: auto;
        overflow-y: hidden;
        white-space: nowrap;
        padding: 5px;
        background-color: #ffffff;
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    .stTabs [data-baseweb="tab-list"]::-webkit-scrollbar {
        height: 6px;
    }
    
    .stTabs [data-baseweb="tab-list"]::-webkit-scrollbar-track {
        background: #f1f1f1;
        border-radius: 10px;
    }
    
    .stTabs [data-baseweb="tab-list"]::-webkit-scrollbar-thumb {
        background: #888;
        border-radius: 10px;
    }
    
    .stTabs [data-baseweb="tab-list"]::-webkit-scrollbar-thumb:hover {
        background: #555;
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 45px;
        padding: 8px 16px;
        background-color: #f0f2f6;
        border-radius: 6px;
        font-size: 0.9rem;
        font-weight: 500;
        color: #444;
        border: 1px solid #e0e0e0;
        transition: all 0.3s ease;
        white-space: nowrap;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        min-width: auto;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background-color: #e8eaf6;
        border-color: #1f77b4;
        transform: translateY(-2px);
        box-shadow: 0 2px 8px rgba(31, 119, 180, 0.2);
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #1f77b4 0%, #1565c0 100%);
        color: white !important;
        border-color: #1565c0;
        box-shadow: 0 4px 12px rgba(31, 119, 180, 0.3);
        font-weight: 600;
    }
    
    /* Tab panel content */
    .stTabs [data-baseweb="tab-panel"] {
        padding: 20px 10px;
    }
    
    /* Input fields styling */
    div[data-testid="stNumberInput"] input {
        font-weight: bold;
        color: #1f77b4;
        border-radius: 6px;
    }
    
    /* Button styling */
    .stButton button {
        border-radius: 6px;
        font-weight: 500;
        transition: all 0.3s ease;
    }
    
    .stButton button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        font-weight: 600;
        color: #1f77b4;
        border-radius: 6px;
    }
    
    /* Dataframe styling */
    .dataframe {
        border-radius: 8px;
        overflow: hidden;
    }
    
    /* Sidebar styling */
    section[data-testid="stSidebar"] {
        background-color: #f8f9fa;
    }
    
    section[data-testid="stSidebar"] .block-container {
        padding-top: 2rem;
    }
    
    /* Info/Warning/Error boxes */
    .stAlert {
        border-radius: 8px;
        border-left: 4px solid;
    }
    
    /* Responsive design for smaller screens */
    @media (max-width: 768px) {
        .stTabs [data-baseweb="tab"] {
            font-size: 0.8rem;
            padding: 6px 12px;
            height: 40px;
        }
        
        .main-header {
            font-size: 1.8rem;
        }
    }
</style>
""", unsafe_allow_html=True)

# Khởi tạo session state
if 'data_extracted' not in st.session_state:
    st.session_state.data_extracted = False
if 'customer_info' not in st.session_state:
    st.session_state.customer_info = {}
if 'financial_info' not in st.session_state:
    st.session_state.financial_info = {}
if 'collateral_info' not in st.session_state:
    st.session_state.collateral_info = {}
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'data_modified' not in st.session_state:
    st.session_state.data_modified = False
if 'uploaded_content' not in st.session_state:
    st.session_state.uploaded_content = ""
if 'last_request_time' not in st.session_state:
    st.session_state.last_request_time = 0

# Hàm định dạng số
def format_number(num):
    """Định dạng số với dấu chấm phân cách hàng nghìn"""
    try:
        return "{:,.0f}".format(float(num)).replace(",", ".")
    except:
        return str(num)

def parse_number(text):
    """Chuyển đổi text thành số"""
    try:
        clean_text = str(text).replace(".", "").replace(",", ".")
        return float(clean_text)
    except:
        return 0

# Hàm trích xuất thông tin từ file docx
def extract_info_from_docx(file):
    """Trích xuất thông tin từ file docx"""
    doc = Document(file)
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    st.session_state.uploaded_content = full_text
    
    customer_info = {}
    financial_info = {}
    collateral_info = {}
    
    # Trích xuất thông tin khách hàng
    name_match = re.search(r'Họ và tên:\s*([^\n\r-]+)', full_text)
    if name_match:
        customer_info['name'] = name_match.group(1).strip()
    
    cccd_match = re.search(r'(?:CMND/)?CCCD(?:/hộ chiếu)?:\s*(\d+)', full_text)
    if cccd_match:
        customer_info['cccd'] = cccd_match.group(1).strip()
    
    address_match = re.search(r'Nơi cư trú:\s*([^\n\r]+)', full_text)
    if address_match:
        customer_info['address'] = address_match.group(1).strip()
    
    phone_match = re.search(r'Số điện thoại:\s*(\d+)', full_text)
    if phone_match:
        customer_info['phone'] = phone_match.group(1).strip()
    
    email_match = re.search(r'Email:\s*([^\s\n\r]+)', full_text)
    if email_match:
        customer_info['email'] = email_match.group(1).strip()
    
    # Trích xuất thông tin tài chính
    total_need_match = re.search(r'Tổng nhu cầu vốn:\s*([\d.,]+)\s*đồng', full_text)
    if total_need_match:
        financial_info['total_need'] = parse_number(total_need_match.group(1))
    
    equity_match = re.search(r'Vốn đối ứng[^:]*:\s*([\d.,]+)\s*đồng', full_text)
    if equity_match:
        financial_info['equity'] = parse_number(equity_match.group(1))
    
    loan_match = re.search(r'Vốn vay[^:]*số tiền:\s*([\d.,]+)\s*đồng', full_text)
    if loan_match:
        financial_info['loan_amount'] = parse_number(loan_match.group(1))
    
    interest_match = re.search(r'Lãi suất:\s*([\d.,]+)%', full_text)
    if interest_match:
        financial_info['interest_rate'] = float(interest_match.group(1).replace(',', '.'))
    
    term_match = re.search(r'Thời hạn vay:\s*(\d+)\s*tháng', full_text)
    if term_match:
        financial_info['loan_term'] = int(term_match.group(1))
    
    purpose_match = re.search(r'Mục đích vay:\s*([^\n\r]+)', full_text)
    if purpose_match:
        financial_info['purpose'] = purpose_match.group(1).strip()
    
    income_patterns = [
        r'Tổng thu nhập[^:]*:\s*([\d.,]+)\s*đồng',
        r'Thu nhập[^:]*:\s*([\d.,]+)\s*đồng/tháng'
    ]
    for pattern in income_patterns:
        income_match = re.search(pattern, full_text)
        if income_match:
            financial_info['monthly_income'] = parse_number(income_match.group(1))
            break
    
    expense_match = re.search(r'Tổng chi phí hàng tháng:\s*([\d.,]+)', full_text)
    if expense_match:
        financial_info['monthly_expense'] = parse_number(expense_match.group(1))
    
    project_income_match = re.search(r'Thu nhập từ kinh doanh[^:]*:\s*([\d.,]+)\s*đồng/tháng', full_text)
    if project_income_match:
        financial_info['project_income'] = parse_number(project_income_match.group(1))
    
    # Trích xuất thông tin tài sản đảm bảo
    collateral_type_match = re.search(r'Tài sản \d+:\s*([^\n\r.]+)', full_text)
    if collateral_type_match:
        collateral_info['type'] = collateral_type_match.group(1).strip()
    
    collateral_value_patterns = [
        r'Giá trị:\s*([\d.,]+)\s*đồng',
        r'Giá trị[^:]*:\s*([\d.,]+)\s*đồng'
    ]
    for pattern in collateral_value_patterns:
        collateral_value_match = re.search(pattern, full_text)
        if collateral_value_match:
            collateral_info['value'] = parse_number(collateral_value_match.group(1))
            break
    
    collateral_address_match = re.search(r'Địa chỉ:\s*([^\n\r]+?)(?:Diện tích|Giấy|Tỷ lệ|\n|$)', full_text)
    if collateral_address_match:
        collateral_info['address'] = collateral_address_match.group(1).strip()
    
    area_match = re.search(r'Diện tích đất:\s*([\d.,]+)\s*m', full_text)
    if area_match:
        collateral_info['area'] = parse_number(area_match.group(1))
    
    return customer_info, financial_info, collateral_info

# Hàm tính toán các chỉ tiêu tài chính
def calculate_financial_metrics(financial_info):
    """Tính toán các chỉ tiêu tài chính"""
    metrics = {}
    
    loan_amount = financial_info.get('loan_amount', 0)
    interest_rate = financial_info.get('interest_rate', 0)
    loan_term = financial_info.get('loan_term', 0)
    monthly_income = financial_info.get('monthly_income', 0)
    monthly_expense = financial_info.get('monthly_expense', 0)
    project_income = financial_info.get('project_income', 0)
    
    if loan_amount > 0 and interest_rate > 0 and loan_term > 0:
        monthly_rate = (interest_rate / 100) / 12
        
        if monthly_rate > 0:
            monthly_payment = loan_amount * (monthly_rate * (1 + monthly_rate)**loan_term) / \
                            ((1 + monthly_rate)**loan_term - 1)
        else:
            monthly_payment = loan_amount / loan_term
        
        metrics['first_month_payment'] = monthly_payment
        
        total_payment = monthly_payment * loan_term
        total_interest = total_payment - loan_amount
        metrics['total_interest'] = total_interest
        metrics['total_payment'] = total_payment
        
        net_income = monthly_income + project_income - monthly_expense
        metrics['net_income'] = net_income
        
        if monthly_income > 0:
            debt_service_ratio = (monthly_payment / monthly_income) * 100
            metrics['debt_service_ratio'] = debt_service_ratio
        else:
            metrics['debt_service_ratio'] = 0
        
        if monthly_payment > 0:
            dscr = net_income / monthly_payment
            metrics['dscr'] = dscr
        else:
            metrics['dscr'] = 0
        
        surplus = net_income - monthly_payment
        metrics['surplus'] = surplus
    
    return metrics

# Hàm tạo lịch trả nợ
def create_repayment_schedule(loan_amount, interest_rate, loan_term):
    """Tạo lịch trả nợ chi tiết"""
    schedule = []
    monthly_rate = (interest_rate / 100) / 12
    
    if monthly_rate > 0:
        monthly_payment = loan_amount * (monthly_rate * (1 + monthly_rate)**loan_term) / \
                        ((1 + monthly_rate)**loan_term - 1)
    else:
        monthly_payment = loan_amount / loan_term
    
    remaining_balance = loan_amount
    
    for month in range(1, loan_term + 1):
        interest_payment = remaining_balance * monthly_rate
        principal_payment = monthly_payment - interest_payment
        
        if month == loan_term:
            principal_payment = remaining_balance
            monthly_payment = principal_payment + interest_payment
        
        schedule.append({
            'Kỳ': month,
            'Dư nợ đầu kỳ': remaining_balance,
            'Tiền gốc': principal_payment,
            'Tiền lãi': interest_payment,
            'Tổng trả': monthly_payment,
            'Dư nợ cuối kỳ': remaining_balance - principal_payment
        })
        
        remaining_balance -= principal_payment
    
    return pd.DataFrame(schedule)

# Hàm xuất Excel
def export_to_excel(df):
    """Xuất DataFrame sang Excel"""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Lịch Trả Nợ')
        
        workbook = writer.book
        worksheet = writer.sheets['Lịch Trả Nợ']
        
        for column in worksheet.columns:
            max_length = 0
            column = [cell for cell in column]
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = (max_length + 2)
            worksheet.column_dimensions[column[0].column_letter].width = adjusted_width
    
    return output.getvalue()

# Hàm xuất báo cáo Word
def export_appraisal_report(customer_info, financial_info, collateral_info, metrics, analysis_file="", analysis_metrics=""):
    """Xuất báo cáo thẩm định sang Word"""
    doc = Document()
    
    title = doc.add_heading('BÁO CÁO THẨM ĐỊNH PHƯƠNG ÁN KINH DOANH', 0)
    title.alignment = 1
    
    doc.add_heading('I. THÔNG TIN KHÁCH HÀNG', 1)
    doc.add_paragraph(f"Họ và tên: {customer_info.get('name', 'N/A')}")
    doc.add_paragraph(f"CCCD: {customer_info.get('cccd', 'N/A')}")
    doc.add_paragraph(f"Địa chỉ: {customer_info.get('address', 'N/A')}")
    doc.add_paragraph(f"Điện thoại: {customer_info.get('phone', 'N/A')}")
    doc.add_paragraph(f"Email: {customer_info.get('email', 'N/A')}")
    
    doc.add_heading('II. THÔNG TIN TÀI CHÍNH', 1)
    doc.add_paragraph(f"Số tiền vay: {format_number(financial_info.get('loan_amount', 0))} đồng")
    doc.add_paragraph(f"Lãi suất: {financial_info.get('interest_rate', 0)}%/năm")
    doc.add_paragraph(f"Thời hạn: {financial_info.get('loan_term', 0)} tháng")
    doc.add_paragraph(f"Mục đích vay: {financial_info.get('purpose', 'N/A')}")
    doc.add_paragraph(f"Thu nhập hàng tháng: {format_number(financial_info.get('monthly_income', 0))} đồng")
    doc.add_paragraph(f"Chi phí hàng tháng: {format_number(financial_info.get('monthly_expense', 0))} đồng")
    
    doc.add_heading('III. TÀI SẢN ĐẢM BẢO', 1)
    doc.add_paragraph(f"Loại tài sản: {collateral_info.get('type', 'N/A')}")
    doc.add_paragraph(f"Giá trị: {format_number(collateral_info.get('value', 0))} đồng")
    doc.add_paragraph(f"Địa chỉ: {collateral_info.get('address', 'N/A')}")
    
    doc.add_heading('IV. CÁC CHỈ TIÊU TÀI CHÍNH', 1)
    doc.add_paragraph(f"Trả nợ hàng tháng: {format_number(metrics.get('first_month_payment', 0))} đồng")
    doc.add_paragraph(f"Thu nhập ròng: {format_number(metrics.get('net_income', 0))} đồng")
    doc.add_paragraph(f"Tỷ lệ trả nợ/thu nhập: {metrics.get('debt_service_ratio', 0):.2f}%")
    doc.add_paragraph(f"DSCR: {metrics.get('dscr', 0):.2f}")
    doc.add_paragraph(f"Số dư sau trả nợ: {format_number(metrics.get('surplus', 0))} đồng")
    doc.add_paragraph(f"Tổng lãi phải trả: {format_number(metrics.get('total_interest', 0))} đồng")
    
    if analysis_file or analysis_metrics:
        doc.add_heading('V. PHÂN TÍCH AI', 1)
        if analysis_file:
            doc.add_heading('Phân tích từ file gốc:', 2)
            doc.add_paragraph(analysis_file)
        if analysis_metrics:
            doc.add_heading('Phân tích từ các chỉ số tài chính:', 2)
            doc.add_paragraph(analysis_metrics)
    
    doc.add_paragraph(f"\nNgày lập: {datetime.now().strftime('%d/%m/%Y')}")
    
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# Hàm cấu hình Gemini
def configure_gemini(api_key):
    """Cấu hình Gemini API"""
    if GENAI_AVAILABLE:
        genai.configure(api_key=api_key)

# Hàm retry với exponential backoff
def retry_with_backoff(func, max_retries=3, base_delay=1):
    """Retry function với exponential backoff"""
    for attempt in range(max_retries):
        try:
            return func()
        except Exception as e:
            if attempt == max_retries - 1:
                raise e
            delay = base_delay * (2 ** attempt)
            time.sleep(delay)

# Hàm phân tích với Gemini
def analyze_with_gemini(api_key, analysis_type, content):
    """Phân tích dữ liệu bằng Gemini AI"""
    if not GENAI_AVAILABLE:
        return "⚠️ Thư viện google-generativeai chưa được cài đặt!"
    
    try:
        configure_gemini(api_key)
        
        def make_request():
            model = genai.GenerativeModel('gemini-2.0-flash')
            
            if analysis_type == "file":
                prompt = f"""
Bạn là chuyên gia thẩm định tín dụng ngân hàng. Hãy phân tích phương án kinh doanh sau và đưa ra đánh giá chi tiết:

{content}

Vui lòng phân tích theo các khía cạnh:
1. Tính khả thi của dự án
2. Khả năng tài chính của khách hàng
3. Rủi ro tiềm ẩn
4. Khuyến nghị cho ngân hàng (nên cho vay hay từ chối, điều kiện gì)
"""
            else:
                prompt = f"""
Bạn là chuyên gia thẩm định tín dụng. Dựa trên các chỉ số tài chính sau, hãy đánh giá khả năng trả nợ và rủi ro:

{content}

Hãy phân tích:
1. Đánh giá các chỉ số tài chính (DSCR, DTI, LTV)
2. Khả năng trả nợ
3. Mức độ rủi ro
4. Khuyến nghị cuối cùng
"""
            
            response = model.generate_content(prompt)
            return response.text
        
        return retry_with_backoff(make_request)
        
    except Exception as e:
        return f"❌ Lỗi khi phân tích: {str(e)}\n\nVui lòng kiểm tra API key và kết nối internet."

# ===== GIAO DIỆN CHÍNH =====

# Header
st.markdown('<div class="main-header">🏦 HỆ THỐNG THẨM ĐỊNH PHƯƠNG ÁN KINH DOANH</div>', unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.header("⚙️ CẤU HÌNH")
    
    uploaded_file = st.file_uploader(
        "📤 Upload file PASDV (.docx)",
        type=['docx'],
        help="Chọn file phương án sử dụng vốn định dạng .docx"
    )
    
    if uploaded_file is not None:
        if st.button("🔍 Trích Xuất Dữ Liệu", use_container_width=True):
            with st.spinner("Đang xử lý..."):
                customer_info, financial_info, collateral_info = extract_info_from_docx(uploaded_file)
                st.session_state.customer_info = customer_info
                st.session_state.financial_info = financial_info
                st.session_state.collateral_info = collateral_info
                st.session_state.data_extracted = True
                st.session_state.data_modified = False
                st.success("✅ Trích xuất thành công!")
                st.rerun()
    
    st.markdown("---")
    
    st.subheader("🤖 AI Configuration")
    api_key = st.text_input(
        "Gemini API Key",
        type="password",
        help="Nhập API key từ Google AI Studio"
    )
    
    if api_key:
        st.success("✅ API Key đã được cấu hình")
    else:
        st.info("💡 Cần API key để sử dụng tính năng AI")
    
    st.markdown("---")
    
    if st.session_state.data_extracted:
        st.subheader("📊 Trạng Thái Dữ Liệu")
        st.success("✅ Dữ liệu đã được tải")
        
        if st.session_state.data_modified:
            st.warning("⚠️ Dữ liệu đã được chỉnh sửa")
        
        if st.button("🔄 Reset Dữ Liệu", use_container_width=True):
            st.session_state.data_extracted = False
            st.session_state.customer_info = {}
            st.session_state.financial_info = {}
            st.session_state.collateral_info = {}
            st.session_state.data_modified = False
            st.rerun()

# Main content
if st.session_state.data_extracted:
    
    # Tạo tabs với tên ngắn gọn và icon
    tabs = st.tabs([
        "👤 Khách hàng",
        "💰 Tài chính",
        "🏠 Tài sản",
        "📊 Chỉ tiêu",
        "📅 Lịch trả nợ",
        "🤖 AI File",
        "🤖 AI Metrics",
        "💬 Chatbot",
        "📥 Xuất file"
    ])
    
    # TAB 1: Thông tin khách hàng
    with tabs[0]:
        st.subheader("👤 Thông Tin Khách Hàng")
        
        col1, col2 = st.columns(2)
        
        with col1:
            name = st.text_input("Họ và tên", value=st.session_state.customer_info.get('name', ''))
            cccd = st.text_input("CCCD", value=st.session_state.customer_info.get('cccd', ''))
            phone = st.text_input("Số điện thoại", value=st.session_state.customer_info.get('phone', ''))
        
        with col2:
            email = st.text_input("Email", value=st.session_state.customer_info.get('email', ''))
            address = st.text_area("Địa chỉ", value=st.session_state.customer_info.get('address', ''), height=100)
        
        if st.button("💾 Lưu Thông Tin Khách Hàng", use_container_width=True):
            st.session_state.customer_info.update({
                'name': name,
                'cccd': cccd,
                'phone': phone,
                'email': email,
                'address': address
            })
            st.session_state.data_modified = True
            st.success("✅ Đã lưu thông tin khách hàng!")
    
    # TAB 2: Thông tin tài chính
    with tabs[1]:
        st.subheader("💰 Thông Tin Tài Chính")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("**Nhu cầu vốn**")
            total_need = st.number_input(
                "Tổng nhu cầu (đồng)",
                value=float(st.session_state.financial_info.get('total_need', 0)),
                step=1000000.0,
                format="%.0f"
            )
            equity = st.number_input(
                "Vốn đối ứng (đồng)",
                value=float(st.session_state.financial_info.get('equity', 0)),
                step=1000000.0,
                format="%.0f"
            )
            loan_amount = st.number_input(
                "Số tiền vay (đồng)",
                value=float(st.session_state.financial_info.get('loan_amount', 0)),
                step=1000000.0,
                format="%.0f"
            )
        
        with col2:
            st.markdown("**Điều kiện vay**")
            interest_rate = st.number_input(
                "Lãi suất (%/năm)",
                value=float(st.session_state.financial_info.get('interest_rate', 0)),
                step=0.1,
                format="%.2f"
            )
            loan_term = st.number_input(
                "Thời hạn (tháng)",
                value=int(st.session_state.financial_info.get('loan_term', 0)),
                step=1
            )
            purpose = st.text_input(
                "Mục đích vay",
                value=st.session_state.financial_info.get('purpose', '')
            )
        
        with col3:
            st.markdown("**Thu chi hàng tháng**")
            monthly_income = st.number_input(
                "Thu nhập (đồng/tháng)",
                value=float(st.session_state.financial_info.get('monthly_income', 0)),
                step=1000000.0,
                format="%.0f"
            )
            monthly_expense = st.number_input(
                "Chi phí (đồng/tháng)",
                value=float(st.session_state.financial_info.get('monthly_expense', 0)),
                step=1000000.0,
                format="%.0f"
            )
            project_income = st.number_input(
                "Thu từ dự án (đồng/tháng)",
                value=float(st.session_state.financial_info.get('project_income', 0)),
                step=1000000.0,
                format="%.0f"
            )
        
        if st.button("💾 Lưu Thông Tin Tài Chính", use_container_width=True):
            st.session_state.financial_info.update({
                'total_need': total_need,
                'equity': equity,
                'loan_amount': loan_amount,
                'interest_rate': interest_rate,
                'loan_term': loan_term,
                'purpose': purpose,
                'monthly_income': monthly_income,
                'monthly_expense': monthly_expense,
                'project_income': project_income
            })
            st.session_state.data_modified = True
            st.success("✅ Đã lưu thông tin tài chính!")
    
    # TAB 3: Tài sản đảm bảo
    with tabs[2]:
        st.subheader("🏠 Tài Sản Đảm Bảo")
        
        col1, col2 = st.columns(2)
        
        with col1:
            collateral_type = st.text_input(
                "Loại tài sản",
                value=st.session_state.collateral_info.get('type', '')
            )
            collateral_value = st.number_input(
                "Giá trị (đồng)",
                value=float(st.session_state.collateral_info.get('value', 0)),
                step=1000000.0,
                format="%.0f"
            )
        
        with col2:
            collateral_address = st.text_area(
                "Địa chỉ tài sản",
                value=st.session_state.collateral_info.get('address', ''),
                height=100
            )
            area = st.number_input(
                "Diện tích (m²)",
                value=float(st.session_state.collateral_info.get('area', 0)),
                step=1.0,
                format="%.2f"
            )
        
        if st.button("💾 Lưu Thông Tin Tài Sản", use_container_width=True):
            st.session_state.collateral_info.update({
                'type': collateral_type,
                'value': collateral_value,
                'address': collateral_address,
                'area': area
            })
            st.session_state.data_modified = True
            st.success("✅ Đã lưu thông tin tài sản!")
    
    # TAB 4: Các chỉ tiêu tài chính
    with tabs[3]:
        st.subheader("📊 Các Chỉ Tiêu Tài Chính")
        
        if st.button("🔄 Tính Toán Lại", use_container_width=True):
            metrics = calculate_financial_metrics(st.session_state.financial_info)
            st.session_state.metrics = metrics
            st.success("✅ Đã tính toán xong!")
        
        if 'metrics' in st.session_state:
            metrics = st.session_state.metrics
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric(
                    "Trả nợ hàng tháng",
                    f"{format_number(metrics.get('first_month_payment', 0))} đ"
                )
            with col2:
                st.metric(
                    "Thu nhập ròng",
                    f"{format_number(metrics.get('net_income', 0))} đ"
                )
            with col3:
                dti = metrics.get('debt_service_ratio', 0)
                st.metric(
                    "Tỷ lệ DTI",
                    f"{dti:.2f}%",
                    delta="Tốt" if dti < 40 else "Cao"
                )
            with col4:
                dscr = metrics.get('dscr', 0)
                st.metric(
                    "DSCR",
                    f"{dscr:.2f}",
                    delta="Tốt" if dscr > 1.25 else "Thấp"
                )
            
            st.markdown("---")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### 💵 Chi Tiết Tài Chính")
                st.write(f"**Số dư sau trả nợ:** {format_number(metrics.get('surplus', 0))} đồng")
                st.write(f"**Tổng lãi phải trả:** {format_number(metrics.get('total_interest', 0))} đồng")
                st.write(f"**Tổng thanh toán:** {format_number(metrics.get('total_payment', 0))} đồng")
            
            with col2:
                st.markdown("### 🎯 Đánh Giá")
                
                if dscr >= 1.25:
                    st.success("✅ DSCR tốt - Khả năng trả nợ cao")
                elif dscr >= 1.0:
                    st.warning("⚠️ DSCR chấp nhận được - Cần theo dõi")
                else:
                    st.error("❌ DSCR thấp - Rủi ro cao")
                
                if dti < 40:
                    st.success("✅ DTI tốt - Gánh nặng nợ hợp lý")
                elif dti < 50:
                    st.warning("⚠️ DTI cao - Cần cân nhắc")
                else:
                    st.error("❌ DTI quá cao - Rủi ro lớn")
                
                ltv = (st.session_state.financial_info.get('loan_amount', 0) / 
                       st.session_state.collateral_info.get('value', 1)) * 100
                st.write(f"**LTV:** {ltv:.2f}%")
                
                if ltv < 70:
                    st.success("✅ LTV tốt")
                elif ltv < 80:
                    st.warning("⚠️ LTV trung bình")
                else:
                    st.error("❌ LTV cao")
            
            # Biểu đồ phân tích
            if PLOTLY_AVAILABLE:
                st.markdown("---")
                st.markdown("### 📊 Biểu Đồ Phân Tích")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    # Biểu đồ cơ cấu trả gốc lãi tháng đầu
                    loan_amount = st.session_state.financial_info.get('loan_amount', 0)
                    interest_rate = st.session_state.financial_info.get('interest_rate', 0)
                    monthly_rate = (interest_rate / 100) / 12
                    
                    interest_first_month = loan_amount * monthly_rate
                    principal_first_month = metrics.get('first_month_payment', 0) - interest_first_month
                    
                    fig_pie = go.Figure(data=[go.Pie(
                        labels=['Tiền gốc', 'Tiền lãi'],
                        values=[principal_first_month, interest_first_month],
                        marker=dict(colors=['#1f77b4', '#ff7f0e']),
                        textinfo='label+percent+value',
                        texttemplate='<b>%{label}</b><br>%{percent}<br>%{value:,.0f} đ',
                        hovertemplate='<b>%{label}</b><br>Số tiền: %{value:,.0f} đồng<br>Tỷ lệ: %{percent}<extra></extra>'
                    )])
                    
                    fig_pie.update_layout(
                        title='Cơ Cấu Trả Nợ Tháng Đầu',
                        height=400,
                        showlegend=True
                    )
                    
                    st.plotly_chart(fig_pie, use_container_width=True)
                
                with col2:
                    # Biểu đồ thu nhập, chi phí, trả nợ
                    monthly_income = st.session_state.financial_info.get('monthly_income', 0)
                    project_income = st.session_state.financial_info.get('project_income', 0)
                    monthly_expense = st.session_state.financial_info.get('monthly_expense', 0)
                    monthly_payment = metrics.get('first_month_payment', 0)
                    
                    total_income = monthly_income + project_income
                    surplus = metrics.get('surplus', 0)
                    
                    fig_bar = go.Figure()
                    
                    fig_bar.add_trace(go.Bar(
                        name='Thu nhập',
                        x=['Tài chính hàng tháng'],
                        y=[total_income],
                        marker_color='#2ecc71',
                        text=[f'{format_number(total_income)}'],
                        textposition='outside',
                        hovertemplate='<b>Thu nhập</b><br>%{y:,.0f} đồng<extra></extra>'
                    ))
                    
                    fig_bar.add_trace(go.Bar(
                        name='Chi phí sinh hoạt',
                        x=['Tài chính hàng tháng'],
                        y=[monthly_expense],
                        marker_color='#e74c3c',
                        text=[f'{format_number(monthly_expense)}'],
                        textposition='outside',
                        hovertemplate='<b>Chi phí sinh hoạt</b><br>%{y:,.0f} đồng<extra></extra>'
                    ))
                    
                    fig_bar.add_trace(go.Bar(
                        name='Trả nợ hàng tháng',
                        x=['Tài chính hàng tháng'],
                        y=[monthly_payment],
                        marker_color='#f39c12',
                        text=[f'{format_number(monthly_payment)}'],
                        textposition='outside',
                        hovertemplate='<b>Trả nợ hàng tháng</b><br>%{y:,.0f} đồng<extra></extra>'
                    ))
                    
                    fig_bar.add_trace(go.Bar(
                        name='Số dư sau trả nợ',
                        x=['Tài chính hàng tháng'],
                        y=[surplus],
                        marker_color='#3498db' if surplus > 0 else '#e74c3c',
                        text=[f'{format_number(surplus)}'],
                        textposition='outside',
                        hovertemplate='<b>Số dư</b><br>%{y:,.0f} đồng<extra></extra>'
                    ))
                    
                    fig_bar.update_layout(
                        title='Thu Nhập, Chi Phí & Trả Nợ Hàng Tháng',
                        yaxis_title='Số tiền (đồng)',
                        height=400,
                        barmode='group',
                        showlegend=True,
                        legend=dict(
                            orientation="h",
                            yanchor="bottom",
                            y=1.02,
                            xanchor="right",
                            x=1
                        )
                    )
                    
                    st.plotly_chart(fig_bar, use_container_width=True)
    
    # TAB 5: Lịch trả nợ
    with tabs[4]:
        st.subheader("📅 Lịch Trả Nợ Chi Tiết")
        
        if st.button("📊 Tạo Lịch Trả Nợ", use_container_width=True):
            loan_amount = st.session_state.financial_info.get('loan_amount', 0)
            interest_rate = st.session_state.financial_info.get('interest_rate', 0)
            loan_term = st.session_state.financial_info.get('loan_term', 0)
            
            if loan_amount > 0 and loan_term > 0:
                schedule = create_repayment_schedule(loan_amount, interest_rate, loan_term)
                st.session_state.repayment_schedule = schedule
                st.success("✅ Đã tạo lịch trả nợ!")
            else:
                st.error("❌ Vui lòng nhập đầy đủ thông tin vay!")
        
        if 'repayment_schedule' in st.session_state:
            df = st.session_state.repayment_schedule.copy()
            
            for col in ['Dư nợ đầu kỳ', 'Tiền gốc', 'Tiền lãi', 'Tổng trả', 'Dư nợ cuối kỳ']:
                df[col] = df[col].apply(format_number)
            
            st.dataframe(df, use_container_width=True, height=400)
            
            if PLOTLY_AVAILABLE:
                st.markdown("### 📈 Biểu Đồ Trả Nợ")
                
                fig = go.Figure()
                fig.add_trace(go.Bar(
                    name='Tiền gốc',
                    x=st.session_state.repayment_schedule['Kỳ'],
                    y=st.session_state.repayment_schedule['Tiền gốc'],
                    marker_color='#1f77b4'
                ))
                fig.add_trace(go.Bar(
                    name='Tiền lãi',
                    x=st.session_state.repayment_schedule['Kỳ'],
                    y=st.session_state.repayment_schedule['Tiền lãi'],
                    marker_color='#ff7f0e'
                ))
                
                fig.update_layout(
                    barmode='stack',
                    title='Cơ Cấu Trả Nợ Theo Kỳ',
                    xaxis_title='Kỳ',
                    yaxis_title='Số tiền (đồng)',
                    height=400
                )
                
                st.plotly_chart(fig, use_container_width=True)
    
    # TAB 6: Phân tích AI (File)
    with tabs[5]:
        st.subheader("🤖 Phân Tích AI - File Gốc")
        
        if not api_key:
            st.warning("⚠️ Vui lòng nhập API Key ở sidebar!")
        elif not GENAI_AVAILABLE:
            st.error("⚠️ Thư viện google-generativeai chưa được cài đặt!")
        else:
            if st.session_state.uploaded_content:
                if st.button("🔍 Phân Tích File Gốc", use_container_width=True):
                    with st.spinner("Đang phân tích..."):
                        analysis = analyze_with_gemini(api_key, "file", st.session_state.uploaded_content)
                        st.session_state.analysis_file = analysis
                
                if 'analysis_file' in st.session_state:
                    st.markdown("#### Kết Quả Phân Tích:")
                    st.info(f"**Nguồn dữ liệu:** File gốc đã upload")
                    st.write(st.session_state.analysis_file)
            else:
                st.warning("⚠️ Chưa có nội dung file để phân tích!")
    
    # TAB 7: Phân tích AI (Metrics)
    with tabs[6]:
        st.subheader("🤖 Phân Tích AI - Chỉ Số Tài Chính")
        
        if not api_key:
            st.warning("⚠️ Vui lòng nhập API Key ở sidebar!")
        elif not GENAI_AVAILABLE:
            st.error("⚠️ Thư viện google-generativeai chưa được cài đặt!")
        else:
            if 'metrics' in st.session_state:
                if st.button("🔍 Phân Tích Chỉ Số", use_container_width=True):
                    data_content = f"""
THÔNG TIN KHÁCH HÀNG:
- Tên: {st.session_state.customer_info.get('name', 'N/A')}
- CCCD: {st.session_state.customer_info.get('cccd', 'N/A')}

THÔNG TIN THU NHẬP:
- Thu nhập hàng tháng: {format_number(st.session_state.financial_info.get('monthly_income', 0))} đồng
- Chi phí hàng tháng: {format_number(st.session_state.financial_info.get('monthly_expense', 0))} đồng

THÔNG TIN VAY VỐN:
- Số tiền vay: {format_number(st.session_state.financial_info.get('loan_amount', 0))} đồng
- Lãi suất: {st.session_state.financial_info.get('interest_rate', 0)}%/năm
- Thời hạn: {st.session_state.financial_info.get('loan_term', 0)} tháng

CÁC CHỈ TIÊU TÀI CHÍNH:
- Trả nợ hàng tháng: {format_number(st.session_state.metrics.get('first_month_payment', 0))} đồng
- Thu nhập ròng: {format_number(st.session_state.metrics.get('net_income', 0))} đồng
- Tỷ lệ trả nợ/thu nhập: {st.session_state.metrics.get('debt_service_ratio', 0):.2f}%
- DSCR: {st.session_state.metrics.get('dscr', 0):.2f}
- Số dư sau trả nợ: {format_number(st.session_state.metrics.get('surplus', 0))} đồng
- Tổng lãi phải trả: {format_number(st.session_state.metrics.get('total_interest', 0))} đồng

TÀI SẢN ĐẢM BẢO:
- Loại: {st.session_state.collateral_info.get('type', 'N/A')}
- Giá trị: {format_number(st.session_state.collateral_info.get('value', 0))} đồng
- LTV: {(st.session_state.financial_info.get('loan_amount', 0) / st.session_state.collateral_info.get('value', 1) * 100):.2f}%
"""
                    with st.spinner("Đang phân tích..."):
                        analysis = analyze_with_gemini(api_key, "metrics", data_content)
                        st.session_state.analysis_metrics = analysis
                
                if 'analysis_metrics' in st.session_state:
                    st.markdown("#### Kết Quả Phân Tích:")
                    st.info(f"**Nguồn dữ liệu:** Các chỉ số tài chính đã nhập")
                    st.write(st.session_state.analysis_metrics)
    
    # TAB 8: Chatbox AI
    with tabs[7]:
        st.subheader("💬 Chatbot AI Gemini")
        
        if not api_key:
            st.warning("⚠️ Vui lòng nhập API Key ở sidebar!")
        elif not GENAI_AVAILABLE:
            st.error("⚠️ Thư viện google-generativeai chưa được cài đặt!")
        else:
            chat_container = st.container()
            with chat_container:
                for i, chat in enumerate(st.session_state.chat_history):
                    if chat['role'] == 'user':
                        st.markdown(f"**👤 Bạn:** {chat['content']}")
                    else:
                        st.markdown(f"**🤖 AI:** {chat['content']}")
                    st.markdown("---")
            
            col1, col2 = st.columns([5, 1])
            with col1:
                user_input = st.text_input("Nhập câu hỏi:", key="chat_input")
            with col2:
                if st.button("Gửi", use_container_width=True):
                    if user_input:
                        st.session_state.chat_history.append({
                            'role': 'user',
                            'content': user_input
                        })
                        
                        context = f"""
Thông tin khách hàng và dự án:
- Tên: {st.session_state.customer_info.get('name', 'N/A')}
- Số tiền vay: {format_number(st.session_state.financial_info.get('loan_amount', 0))} đồng
- Lãi suất: {st.session_state.financial_info.get('interest_rate', 0)}%
- Thu nhập: {format_number(st.session_state.financial_info.get('monthly_income', 0))} đồng/tháng
"""
                        
                        with st.spinner("AI đang suy nghĩ..."):
                            try:
                                current_time = time.time()
                                time_since_last = current_time - st.session_state.last_request_time
                                if time_since_last < 2:
                                    time.sleep(2 - time_since_last)
                                
                                configure_gemini(api_key)
                                
                                def chat_request():
                                    model = genai.GenerativeModel('gemini-2.0-flash')
                                    prompt = f"{context}\n\nCâu hỏi: {user_input}"
                                    response = model.generate_content(prompt)
                                    st.session_state.last_request_time = time.time()
                                    return response.text
                                
                                ai_response = retry_with_backoff(chat_request)
                                
                                st.session_state.chat_history.append({
                                    'role': 'assistant',
                                    'content': ai_response
                                })
                            except Exception as e:
                                ai_response = f"❌ Lỗi: {str(e)}"
                                st.session_state.chat_history.append({
                                    'role': 'assistant',
                                    'content': ai_response
                                })
                        
                        st.rerun()
            
            if st.button("🗑️ Xóa Lịch Sử", use_container_width=True):
                st.session_state.chat_history = []
                st.rerun()
    
    # TAB 9: Xuất dữ liệu
    with tabs[8]:
        st.subheader("📥 Xuất Dữ Liệu")
        
        export_option = st.selectbox(
            "Chọn loại xuất:",
            ["Bảng trả nợ (Excel)", "Báo cáo thẩm định (Word)"]
        )
        
        if export_option == "Bảng trả nợ (Excel)":
            st.markdown("### 📊 Xuất Bảng Trả Nợ")
            
            if 'repayment_schedule' in st.session_state:
                st.dataframe(st.session_state.repayment_schedule, use_container_width=True)
                
                excel_data = export_to_excel(st.session_state.repayment_schedule)
                st.download_button(
                    label="📥 Tải Xuống Excel",
                    data=excel_data,
                    file_name=f"ke_hoach_tra_no_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            else:
                st.warning("⚠️ Chưa có dữ liệu!")
        
        else:
            st.markdown("### 📄 Xuất Báo Cáo Thẩm Định")
            
            if 'metrics' in st.session_state:
                analysis_file = st.session_state.get('analysis_file', '')
                analysis_metrics = st.session_state.get('analysis_metrics', '')
                
                word_data = export_appraisal_report(
                    st.session_state.customer_info,
                    st.session_state.financial_info,
                    st.session_state.collateral_info,
                    st.session_state.metrics,
                    analysis_file,
                    analysis_metrics
                )
                
                st.download_button(
                    label="📥 Tải Xuống Word",
                    data=word_data,
                    file_name=f"bao_cao_tham_dinh_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.warning("⚠️ Chưa có dữ liệu!")

else:
    st.markdown("""
    <div style='text-align: center; padding: 3rem;'>
        <h2>👋 Chào Mừng Đến Với Hệ Thống Thẩm Định</h2>
        <p style='font-size: 1.2rem; color: #666;'>
            Vui lòng upload file phương án sử dụng vốn (.docx) ở sidebar để bắt đầu!
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        ### 📤 Bước 1: Upload File
        - Click vào sidebar bên trái
        - Chọn file PASDV.docx
        - Click "Trích xuất dữ liệu"
        """)
    
    with col2:
        st.markdown("""
        ### ✏️ Bước 2: Chỉnh Sửa
        - Xem và chỉnh sửa thông tin
        - Sử dụng nút +/- để điều chỉnh
        - Lưu thay đổi khi cần
        """)
    
    with col3:
        st.markdown("""
        ### 📊 Bước 3: Phân Tích
        - Xem các chỉ tiêu tài chính
        - Phân tích bằng AI
        - Xuất báo cáo
        """)
    
    st.markdown("---")
    
    with st.expander("ℹ️ Hướng dẫn lấy Gemini API Key"):
        st.markdown("""
        1. Truy cập: https://aistudio.google.com/app/apikey
        2. Đăng nhập bằng tài khoản Google
        3. Click "Create API Key"
        4. Copy API Key và paste vào ô bên sidebar
        """)

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; padding: 1rem;'>
    <p>🏦 Hệ Thống Thẩm Định Phương Án Kinh Doanh v1.3</p>
    <p>Powered by Streamlit & Google Gemini 2.0 Flash AI</p>
</div>
""", unsafe_allow_html=True)
